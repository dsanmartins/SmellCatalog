## Daniel San Martin
## 19-05-2018
## SAs smell catalog

# import the library
from appJar import gui
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sqlite3
import io
import os
import datetime
conn = sqlite3.connect('smell.db')

##########################################################################################################################################################################
# Functions
##########################################################################################################################################################################
# handle database
def createTables():
    c = conn.cursor()
    # Create table if exists
    c.execute('''CREATE TABLE IF NOT EXISTS smell (Smell text, Description text, rationale text, causes text, exampleDesc text, figure blob, pngpath text, affected text, abstractions text, considerations text)''')
    conn.commit()

def insertSmell():
    c = conn.cursor()
    # Create table if exists
    smellName = app.getEntry("Smell") 
    description = app.getTextArea("Description")
    rationale = app.getTextArea("implRationale") 
    causes = app.getTextArea("implPotential")
    exampleDesc = app.getTextArea("exDesc")
    example = app.getEntry("f1")
    affected = app.getTextArea("implQualityAttr")
    abstractions = app.getTextArea("implAbstrAffec")
    considerations = app.getTextArea("impllblPractialCons")
    
    boolean = searchSmell(smellName)
    if not boolean:
        #Create PNG file
        if example.endswith('png'):
            png = open(example,'rb')
            ablob = png.read()
            smellDocumentation = [(smellName, description, rationale, causes, exampleDesc, sqlite3.Binary(ablob), example, affected, abstractions, considerations)]
            c.executemany('INSERT INTO smell VALUES (?,?,?,?,?,?,?,?,?,?)', smellDocumentation)
            conn.commit()
            app.clearAllEntries(callFunction=False)
            app.clearAllTextAreas(callFunction=False)
            app.infoBox("MsgDone", "Saved!", parent=None)
        elif not example:
            smellDocumentation = [(smellName, description, rationale, causes, exampleDesc, '', example, affected, abstractions, considerations)]
            c.executemany('INSERT INTO smell VALUES (?,?,?,?,?,?,?,?,?,?)', smellDocumentation)
            conn.commit()
            app.clearAllEntries(callFunction=False)
            app.clearAllTextAreas(callFunction=False)
            app.infoBox("MsgDone", "Saved!", parent=None)
        else:
            app.infoBox("MsgInsert", "The file is not a PNG file!", parent=None)
    else:
        question = app.yesNoBox("MsgQuestion", "The smell already exist!. Do you want to overwrite the data?", parent=None)
        if question is True:
            if example.endswith('png'):
                png = open(example,'rb')
                ablob = png.read()
                smellDocumentation = [(description, rationale, causes, exampleDesc, sqlite3.Binary(ablob), example, affected, abstractions, considerations, smellName)]
                c.executemany('UPDATE smell SET Description=?, rationale=?, causes=?, exampleDesc = ?, figure = ?, pngpath=?, affected=?, abstractions=?, considerations = ? WHERE Smell = ?',smellDocumentation)
                conn.commit()
                app.infoBox("MsgUpdate", "Updated!", parent=None)
            elif not example:
                smellDocumentation = [(description, rationale, causes, exampleDesc, '', example, affected, abstractions, considerations, smellName)]
                c.executemany('UPDATE smell SET Description=?, rationale=?, causes=?, exampleDesc = ?, figure = ?, pngpath=?, affected=?, abstractions=?, considerations = ? WHERE Smell = ?', smellDocumentation)
                conn.commit()
                app.infoBox("MsgUpdate", "Updated!", parent=None)
            else:
                app.infoBox("MsgInsert", "The file is not a PNG file!. Data was not updated!.", parent=None)   
             
def searchSmell(op):
    smell = ''
    if op == 1:
        smellName = app.getEntry("Smell") 
        c = conn.cursor()
        c.execute("SELECT * FROM smell WHERE smell = ?",[smellName])
        row = c.fetchone()
        if row == None:
             app.infoBox("MsgEmpty", "The smell does not exist!", parent=None)
        else:
            app.clearAllEntries(callFunction=False)
            app.clearAllTextAreas(callFunction=False)
            smell = row[0]
            description = row[1]
            rationale = row[2]
            causes = row[3]
            exampleDesc = row[4]
            pngpath = row[6]
            affected = row[7]
            abstractions = row[8]
            considerations = row[9]
            app.setEntry("Smell",smell)
            app.setTextArea("Description",description)
            app.setTextArea("implRationale",rationale) 
            app.setTextArea("implPotential",causes)
            app.setTextArea("exDesc",exampleDesc)
            app.setEntry("f1",pngpath)
            app.setTextArea("implQualityAttr",affected)
            app.setTextArea("implAbstrAffec",abstractions)
            app.setTextArea("impllblPractialCons",considerations)
        c.close()
    else:
        c = conn.cursor()
        c.execute("SELECT * FROM smell WHERE smell = ?",[op])
        row = c.fetchone()
        if row == None:
            smell = ''
        else:
            smell = row[0]
    
    if smell:
        return True
    else:
        return False

def deleteSmell():    
    smellName = app.getEntry("Smell")
    c = conn.cursor()
    c.execute("DELETE FROM smell WHERE smell = ?",[smellName])
    app.clearAllEntries(callFunction=False)
    app.clearAllTextAreas(callFunction=False)
    
def getTableHeader():
    c = conn.cursor()
    c.execute("SELECT smell, description FROM smell")
    header =  tuple(list(map(lambda x: x[0], c.description)))
    list1 = []
    list1.append(header)
    return list1

def getAllSmell():
    c = conn.cursor()
    c.execute("SELECT smell, description FROM smell")
    rows = c.fetchall();
    list1 = []
    for row in rows:
        list1.append(row)
    return list1

# handle button events
def press(button):
    if button == "Close":
        conn.close()
        app.stop()
    elif button == "Save":
        str = app.getEntry("Smell") 
        if not str:
            app.infoBox("MsgEmpty", "The smell cannot be empty!", parent=None)
        else:
            insertSmell()
    elif button == "Search":
        str = app.getEntry("Smell") 
        if not str:
            app.infoBox("MsgEmpty", "The smell cannot be empty!", parent=None)
        else:
            searchSmell(1)
    
    elif button == "Delete":
        str = app.getEntry("Smell") 
        if not str:
            app.infoBox("MsgEmpty", "The smell cannot be empty!", parent=None)
        else:
            deleteSmell()

def processRow(pos):
    list1 = app.getTableRow("SmellTable", pos)
    var = list1[0]
    c = conn.cursor()
    c.execute("SELECT * FROM smell WHERE smell = ?",[var])
    row = c.fetchone()
    smell = row[0]
    description = row[1]
    rationale = row[2]
    causes = row[3]
    exDesc = row[4]
    png = row[5]
    affected = row[7]
    abstractions = row[8]
    considerations = row[9]
    
    #Description
    document = Document()
    document.add_heading(smell.title(), 0)
    
    now = datetime.datetime.now()
    dt = now.strftime("%Y-%m-%d %H:%M")
    p = document.add_paragraph()
    run = p.add_run('Generated Automatically on ' + dt).italic = True
    
    p = document.add_paragraph()
    run = p.add_run('Description: ')
    run.bold = True
    run.underline = True
    font = run.font
    font.size = Pt(12)
    p = document.add_paragraph(description)
    paragraph_format = p.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    #Rationale
    p = document.add_paragraph()
    run = p.add_run('Rationale: ')
    run.bold = True
    run.underline = True
    font = run.font
    font.size = Pt(12)
    p = document.add_paragraph(rationale)
    paragraph_format = p.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    #Causes
    p = document.add_paragraph()
    run = p.add_run('Potential Causes: ')
    run.bold = True
    run.underline = True
    font = run.font
    font.size = Pt(12)
    p = document.add_paragraph(causes)
    paragraph_format = p.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    
    #Example and Figure
    p = document.add_paragraph()
    run = p.add_run('Examples: ')
    run.bold = True
    run.underline = True
    font = run.font
    font.size = Pt(12)
    p = document.add_paragraph(exDesc)
    paragraph_format = p.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if png != '':
        if not os.path.exists(smell):
            os.makedirs(smell)
        filename = smell + '/' +smell + '.png'
        with open(filename, 'wb') as output_file:
            output_file.write(png)
        document.add_picture(filename)
    else:
        if not os.path.exists(smell):
            os.makedirs(smell)
    
    #Affected
    p = document.add_paragraph()
    run = p.add_run('Imptacted Quality Attributes: ')
    run.bold = True
    run.underline = True
    font = run.font
    font.size = Pt(12)
    p = document.add_paragraph(affected)
    paragraph_format = p.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    #Abstractions
    p = document.add_paragraph()
    run = p.add_run('Affected Architectural Abstractions: ')
    run.bold = True
    run.underline = True
    font = run.font
    font.size = Pt(12)
    p = document.add_paragraph(abstractions)
    paragraph_format = p.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #Practical Considerations
    p = document.add_paragraph()
    run = p.add_run('Practical Considerations: ')
    run.bold = True
    run.underline = True
    font = run.font
    font.size = Pt(12)
    p = document.add_paragraph(considerations)
    paragraph_format = p.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
  
    document.add_page_break()
    document.save(smell + '/' + smell + '.docx')
    app.infoBox("MsgWord", "A document has been generated successfully!", parent=None)

def onChangeTF(name):
    tab = app.getTabbedFrameSelectedTab(name)
    if tab == 'List of Architectural Smells':
        rows = getAllSmell()
        app.deleteAllTableRows("SmellTable")
        app.addTableRows("SmellTable", rows)

# handle menuPress
def menu(name):
    if name == "About":
        app.infoBox("Information", "Smelly SAs Catalog 2018.\nCreated by Daniel San Martin\ndsanmartins@gmail.com", parent=None)

##########################################################################################################################################################################
# Widgets implementation
##########################################################################################################################################################################

#Call Database
createTables()


# create a GUI variable called app
app = gui("Architectural Smells for SAs Catalog", "1350x768")
app.setBg("black")
app.setFont(18)

############################
#create elements in the grid
############################
#0-0
app.addLabel("lblTitle", "Smelly SAs Catalog", 0, 0, 2)
app.setLabelFg("lblTitle", "white")
app.getLabelWidget("lblTitle").config(font="Verdana 18 italic bold" )

app.startTabbedFrame("TabbedFrame")
app.startTab("Architectural Smell")
app.setBg("black")
app.setFg("white")
app.setTabbedFrameChangeFunction("TabbedFrame", onChangeTF) 


#1-0
app.startFrame("frm0", row=1, column=0)
app.setBg("black")
app.setSticky("w")
app.setStretch("none")

app.addLabel("lblSmell", "Smell Name")
app.setLabelBg("lblSmell", "black")
app.setLabelFg("lblSmell", "white")
app.getLabelWidget("lblSmell").config(font="Verdana 12 italic bold")
app.addEntry('Smell')
app.setEntryMaxLength("Smell", 40)
app.getEntryWidget("Smell").config(width=40)
app.stopFrame()

#1-1
app.startFrame("frm1", row=1, column=1)
app.setBg("black")
app.setSticky("w")
app.setStretch("none")

app.addLabel("lblSmellDesc","A Brief Description of the Smell")
app.setLabelBg("lblSmellDesc", "black")
app.setLabelFg("lblSmellDesc", "white")
app.getLabelWidget("lblSmellDesc").config(font="Verdana 12 italic bold" )
app.addScrolledTextArea('Description')
app.getTextAreaWidget("Description").config(height=4, width= 70)
app.stopFrame()

#2-0
app.startFrame("frm2", row=2, column=0)
app.setBg("black")
app.setSticky("w")
app.setStretch("none")

app.addLabel("lblRationale","Rationale")
app.setLabelBg("lblRationale", "black")
app.setLabelFg("lblRationale", "white")
app.getLabelWidget("lblRationale").config(font="Verdana 12 italic bold" )
app.addScrolledTextArea('implRationale')
app.getTextAreaWidget("implRationale").config(height="4" , width= 70)
app.stopFrame()

#2-1
app.startFrame("frm3", row=2, column=1)
app.setBg("black")
app.setSticky("w")
app.setStretch("none")

app.addLabel("lblPotential","Potential Causes")
app.setLabelBg("lblPotential", "black")
app.setLabelFg("lblPotential", "white")
app.getLabelWidget("lblPotential").config(font="Verdana 12 italic bold" )
app.addScrolledTextArea('implPotential')
app.getTextAreaWidget("implPotential").config(height="4" , width= 70)
app.stopFrame()

#3-0
app.startFrame("frm4", row=3, column=0)
app.setBg("black")
app.setSticky("w")
app.setStretch("none")

app.addLabel("lblExampleDescription","A Description of an Example")
app.setLabelBg("lblExampleDescription", "black")
app.setLabelFg("lblExampleDescription", "white")
app.getLabelWidget("lblExampleDescription").config(font="Verdana 12 italic bold" )
app.addScrolledTextArea("exDesc")
app.getTextAreaWidget("exDesc").config(height=4, width =70)
app.stopFrame()

#3-1
app.startFrame("frm5", row=3, column=1)
app.setBg("black")
app.setSticky("w")
app.setStretch("none")

app.addLabel("lblExamples","A Figure of the Example (PNG file)")
app.setLabelBg("lblExamples", "black")
app.setLabelFg("lblExamples", "white")
app.getLabelWidget("lblExamples").config(font="Verdana 12 italic bold" )
app.addFileEntry("f1")
app.getEntryWidget("f1").config(width =44)
app.stopFrame()

#4-0
app.startFrame("frm6", row=4, column=0)
app.setBg("black")
app.setSticky("w")
app.setStretch("none")

app.addLabel("lblQualityAttr","Impacted Quality Attributes")
app.setLabelBg("lblQualityAttr", "black")
app.setLabelFg("lblQualityAttr", "white")
app.getLabelWidget("lblQualityAttr").config(font="Verdana 12 italic bold" )
app.addScrolledTextArea('implQualityAttr')
app.getTextAreaWidget("implQualityAttr").config(height=4, width= 70)
app.stopFrame()

#4-1
app.startFrame("frm7", row=4, column=1)
app.setBg("black")
#app.setSticky("e")
app.setStretch("none")

app.addLabel("lblAbstrAffec","Affected Architectural Abstractions")
app.setLabelBg("lblAbstrAffec", "black")
app.setLabelFg("lblAbstrAffec", "white")
app.getLabelWidget("lblAbstrAffec").config(font="Verdana 12 italic bold" )
app.addScrolledTextArea('implAbstrAffec')
app.getTextAreaWidget("implAbstrAffec").config(height=4, width= 70)
app.stopFrame()

#5-0
app.startFrame("frm8", row=5, column=0)
app.setBg("black")
#app.setSticky("e")
app.setStretch("none")

app.addLabel("lblPractialCons","Practical Considerations")
app.setLabelBg("lblPractialCons", "black")
app.setLabelFg("lblPractialCons", "white")
app.getLabelWidget("lblPractialCons").config(font="Verdana 12 italic bold" )
app.addScrolledTextArea('impllblPractialCons')
app.getTextAreaWidget("impllblPractialCons").config(height=4, width= 70)
app.stopFrame()

#add menu
fileMenus = ["About"]
app.addMenuList("Menu", fileMenus, menu)

# link the buttons to the function called press
#app.startFrame("frm8", row=8, column=0)
app.addButtons(["Save", "Search", "Delete", "Close"], press,9,0,2)
app.stopTab()

app.startTab("List of Architectural Smells")
app.setBg("black")
app.setFg("white")
header = getTableHeader()
app.addTable("SmellTable", header, action=processRow, actionHeading= "Word")
rows = getAllSmell()
app.addTableRows("SmellTable", rows)
app.stopTab()

app.startTab("Information")
app.setBg("black")
app.setFg("white")
app.addMessage("Name", "Smell & Description: A concise, intuitive name based on our naming scheme (comprises two words: first word is an adjective, and second word is the primarily violated design principle). The name is followed by a concise description of the architectural smell (along with its possible forms).\n--------------------------------------------------------------------------------\nRationale: Reason/justification for the architectural smell in the context of adaptive systems.\n--------------------------------------------------------------------------------\nPotential Causes: List of typical reasons for the occurrence of the smell (a nonexhaustive list based on our experience).\n--------------------------------------------------------------------------------\nExamples: One or more examples highlighting the smell. If a smell has multiple forms, each form may be illustrated using a specific example (Figures in PNG).\n--------------------------------------------------------------------------------\nImpacted Quality Attributes: The quality attributes that are negatively impacted because of this smell. It may include understandability, changeability, extensibility, reusability, testability, reliability among others.\n--------------------------------------------------------------------------------\nAffected Architectural Abstractions: Adaptive System Abstractions affected by the architectural smell.\n--------------------------------------------------------------------------------\nPractical Considerations: Sometimes, in a real-world context, a particular design decision that introduces a smell may be purposely made either due to constraints (such as language or platform limitations) or to address a larger problem in the overall design and implementation.",1,1,1,8)
app.setMessageBg("Name", "black")
app.setMessageFg("Name", "white")
app.getMessageWidget("Name").config(width="1200", font="Verdana 12 italic bold" )
app.stopTab()

app.stopTabbedFrame()
app.setTabbedFrameActiveFg("TabbedFrame", "white")


# start the GUI
app.go()
